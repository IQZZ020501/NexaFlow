from __future__ import annotations

import json
import os
import re
import sys
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


HEADING = re.compile(r"^(#{1,3})\s+(.+)$")
BULLET = re.compile(r"^[-*]\s+(.+)$")
NUMBERED = re.compile(r"^\d+[.)]\s+(.+)$")
INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*.+?\*)")
REPORT_STYLE = "report"
FORMAL_LEGAL_STYLE = "formal_legal"

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
TABLE_USABLE_DXA = 9071  # A4 width 21cm minus 2.5cm side margins
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = 120


def _resolve_cjk_font(*, formal: bool = False) -> str | None:
    """Return the platform CJK family without probing from an isolated child."""
    configured = os.environ.get("NEXAFLOW_CJK_FONT", "").strip()
    if configured:
        return configured
    if formal:
        return {
            "linux": "Noto Serif CJK SC",
            "darwin": "STFangsong",
        }.get(sys.platform)
    return {
        "linux": "Noto Sans CJK SC",
        "darwin": "PingFang SC",
    }.get(sys.platform)


CJK_FONT = _resolve_cjk_font()
FORMAL_CJK_FONT = _resolve_cjk_font(formal=True)


def _display_units(value: str) -> int:
    return sum(
        2 if unicodedata.east_asian_width(character) in {"W", "F", "A"} else 1
        for character in value
    )


def _input() -> tuple[str, str]:
    value = json.load(sys.stdin)
    if not isinstance(value, dict) or not isinstance(value.get("content"), str):
        raise ValueError("documents content must be a string")
    style = value.get("style", REPORT_STYLE)
    if not isinstance(style, str) or style not in {REPORT_STYLE, FORMAL_LEGAL_STYLE}:
        raise ValueError("documents style must be report or formal_legal")
    content = value["content"].strip()
    if not content:
        raise ValueError("documents content is empty")
    if (
        style == FORMAL_LEGAL_STYLE
        and "劳动仲裁不收费" in content
        and re.search(r"仲裁费用[^\n。；]{0,80}由被申请人承担", content)
    ):
        raise ValueError(
            "formal_legal content contains conflicting arbitration fee statements: "
            "'仲裁费用…由被申请人承担' and '劳动仲裁不收费'"
        )
    return content, style


def _set_east_asia_font(target, font: str | None = CJK_FONT) -> None:
    if not font:
        return
    r_pr = target._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), font)


def _set_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    formal: bool = False,
) -> None:
    run.font.name = "Times New Roman" if formal else "Arial"
    _set_east_asia_font(run, FORMAL_CJK_FONT if formal else CJK_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _add_inline(paragraph, text: str, *, formal: bool = False) -> None:
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _set_font(run, bold=True, formal=formal)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            _set_east_asia_font(run, FORMAL_CJK_FONT if formal else CJK_FONT)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            _set_font(run, formal=formal)
            run.italic = True
        else:
            run = paragraph.add_run(part)
            _set_font(run, formal=formal)


def _cells(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_separator(line: str) -> bool:
    cells = _cells(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def _shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _table_column_widths(rows: list[list[str]], count: int) -> list[int]:
    natural = []
    for column_index in range(count):
        units = max(
            _display_units(str(row[column_index]))
            for row in rows
            if column_index < len(row)
        )
        natural.append(min(max(units * 110 + 240, 900), 3600))
    total = sum(natural)
    scale = TABLE_USABLE_DXA / total
    widths = [max(int(value * scale), 1) for value in natural]
    widths[-1] += TABLE_USABLE_DXA - sum(widths)
    return widths


def _apply_table_geometry(table, rows: list[list[str]], count: int) -> None:
    widths = _table_column_widths(rows, count)
    total = sum(widths)
    tbl_pr = table._tbl.tblPr

    def insert(element) -> None:
        lookup = tbl_pr.find(qn("w:tblLook"))
        if lookup is not None:
            lookup.addprevious(element)
        else:
            tbl_pr.append(element)

    table.autofit = False
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        insert(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    for element in tbl_pr.findall(qn("w:tblInd")):
        tbl_pr.remove(element)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    insert(tbl_ind)
    for element in tbl_pr.findall(qn("w:tblLayout")):
        tbl_pr.remove(element)
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    insert(tbl_layout)
    for element in tbl_pr.findall(qn("w:tblCellMar")):
        tbl_pr.remove(element)
    cell_mar = OxmlElement("w:tblCellMar")
    for side in ("top", "bottom", "start", "end"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(CELL_MARGIN_DXA))
        node.set(qn("w:type"), "dxa")
        cell_mar.append(node)
    insert(cell_mar)
    grid = table._tbl.find(qn("w:tblGrid"))
    for grid_col, column_width in zip(grid.findall(qn("w:gridCol")), widths):
        grid_col.set(qn("w:w"), str(column_width))
    for row_index, row in enumerate(table.rows):
        row_properties = row._tr.get_or_add_trPr()
        if row_index == 0:
            row_properties.append(OxmlElement("w:tblHeader"))
        cant_split = OxmlElement("w:cantSplit")
        row_properties.append(cant_split)
        for column_index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                shading = tc_pr.find(qn("w:shd"))
                if shading is not None:
                    shading.addprevious(tc_w)
                else:
                    tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[column_index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_table(
    document: Document,
    rows: list[list[str]],
    *,
    formal: bool = False,
) -> None:
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        for column_index in range(width):
            text = values[column_index] if column_index < len(values) else ""
            cell = table.cell(row_index, column_index)
            cell.text = ""
            _add_inline(cell.paragraphs[0], text, formal=formal)
            if formal:
                cell.paragraphs[0].paragraph_format.first_line_indent = Pt(0)
            if row_index == 0:
                _shade(cell, "E7E6E6" if formal else "D9EAF7")
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    _apply_table_geometry(table, rows, width)


def _add_page_number_footer(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    _set_font(run, size=9)
    run.font.color.rgb = RGBColor(89, 89, 89)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, end):
        run._r.append(element)


def _configure(document: Document, *, formal: bool = False) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    if not formal:
        _add_page_number_footer(section)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman" if formal else "Arial"
    _set_east_asia_font(normal, FORMAL_CJK_FONT if formal else CJK_FONT)
    normal.font.size = Pt(12 if formal else 11)
    normal.paragraph_format.space_after = Pt(0 if formal else 7)
    normal.paragraph_format.line_spacing = 1.4 if formal else 1.15
    normal.paragraph_format.first_line_indent = Pt(24) if formal else None
    normal.paragraph_format.widow_control = True
    for level, size in ((1, 18), (2, 15), (3, 13)):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Arial"
        _set_east_asia_font(style)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(31, 78, 121)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def _render(content: str, output_path: Path, style: str = REPORT_STYLE) -> None:
    formal = style == FORMAL_LEGAL_STYLE
    document = Document()
    _configure(document, formal=formal)
    lines = content.splitlines()
    index = 0
    first_h1 = True
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if formal and line == "---":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            index += 1
            continue
        if (
            "|" in line
            and index + 1 < len(lines)
            and _is_separator(lines[index + 1].strip())
        ):
            rows = [_cells(line)]
            index += 2
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                rows.append(_cells(lines[index]))
                index += 1
            _add_table(document, rows, formal=formal)
            continue
        heading = HEADING.match(line)
        if heading:
            level = len(heading.group(1))
            paragraph = (
                document.add_paragraph()
                if formal
                else document.add_heading(level=level)
            )
            _add_inline(paragraph, heading.group(2), formal=formal)
            if formal:
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.paragraph_format.keep_with_next = True
                paragraph.paragraph_format.keep_together = True
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt({1: 18, 2: 14, 3: 12}[level])
                    run.font.color.rgb = RGBColor(0, 0, 0)
                if level == 1 and first_h1:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    first_h1 = False
        elif match := BULLET.match(line):
            paragraph = document.add_paragraph(
                style="Normal" if formal else "List Bullet"
            )
            _add_inline(paragraph, line if formal else match.group(1), formal=formal)
            if formal:
                paragraph.paragraph_format.first_line_indent = Pt(0)
        elif match := NUMBERED.match(line):
            paragraph = document.add_paragraph(
                style="Normal" if formal else "List Number"
            )
            _add_inline(paragraph, line if formal else match.group(1), formal=formal)
            if formal:
                paragraph.paragraph_format.first_line_indent = Pt(0)
        elif line.startswith("> "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(0)
            _add_inline(paragraph, line[2:], formal=formal)
            if formal:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                paragraph.paragraph_format.left_indent = Cm(0.7)
                for run in paragraph.runs:
                    run.italic = True
        else:
            paragraph = document.add_paragraph()
            _add_inline(paragraph, line, formal=formal)
        index += 1
    document.save(output_path)


def main() -> None:
    output_path = Path(os.environ["NEXAFLOW_OUTPUT_PATH"])
    if output_path.suffix.lower() != ".docx":
        raise ValueError("documents Skill requires a .docx filename")
    content, style = _input()
    _render(content, output_path, style)
    verified = Document(output_path)
    text = "".join(paragraph.text for paragraph in verified.paragraphs).strip()
    table_text = "".join(
        cell.text for table in verified.tables for row in table.rows for cell in row.cells
    ).strip()
    if not text and not table_text:
        raise ValueError("generated DOCX is empty")
    print(
        json.dumps(
            {
                "renderer": "documents",
                "paragraphs": len(verified.paragraphs),
                "tables": len(verified.tables),
            },
            separators=(",", ":"),
        )
    )


if __name__ == "__main__":
    main()
